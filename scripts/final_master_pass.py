from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"C:\Users\arind\Documents\book")
SOURCE_MD = ROOT / "manuscript_part_i_chapter_1.md"
BASE_DOCX = ROOT / "Waking_Up_Together_Complete_Through_Chapter_15.docx"
MASTER_MD = ROOT / "Waking_Up_Together_Master_Manuscript.md"
MASTER_DOCX = ROOT / "Waking_Up_Together_Master_Manuscript.docx"

AUTHOR_PHOTO = ROOT / "assets" / "author-photo.jpeg"

CHAPTER_PAGES = {
    1: 5,
    2: 11,
    3: 15,
    4: 20,
    5: 24,
    6: 29,
    7: 34,
    8: 39,
    9: 44,
    10: 50,
    11: 54,
    12: 59,
    13: 65,
    14: 71,
    15: 77,
}

TOC_LEFT = [
    ("PART I: THE ARCHITECTURE OF THE ILLUSION", None),
    ("Chapter 1: The Code in the Mirror", CHAPTER_PAGES[1]),
    ("Chapter 2: Glitches in the Modern Matrix", CHAPTER_PAGES[2]),
    ("Chapter 3: Breaking the Simulation", CHAPTER_PAGES[3]),
    ("PART II: SUBATOMIC DEVOTION", None),
    ("Chapter 4: The Entanglement Protocol", CHAPTER_PAGES[4]),
    ("Chapter 5: The Chaos Theory of Chemistry", CHAPTER_PAGES[5]),
    ("Chapter 6: The Intellectual Turn-On", CHAPTER_PAGES[6]),
    ("PART III: THE METAPHYSICS OF ABUNDANCE", None),
    ("Chapter 7: Financial Alchemy", CHAPTER_PAGES[7]),
    ("Chapter 8: The Abundance Frequency", CHAPTER_PAGES[8]),
    ("Chapter 9: The Wealth Covenant", CHAPTER_PAGES[9]),
]

TOC_RIGHT = [
    ("PART IV: THE SYMPHONY OF INTIMACY", None),
    ("Chapter 10: The Neuroscience of the Ecstatic", CHAPTER_PAGES[10]),
    ("Chapter 11: Parallel Realities and Choice Paralysis", CHAPTER_PAGES[11]),
    ("Chapter 12: Sacred Geometry in Swipe Culture", CHAPTER_PAGES[12]),
    ("PART V: COSMIC ADVANCEMENT", None),
    ("Chapter 13: Tuning Your Frequency", CHAPTER_PAGES[13]),
    ("Chapter 14: The Architecture of the Power Couple", CHAPTER_PAGES[14]),
    ("Chapter 15: The Final Code", CHAPTER_PAGES[15]),
]

TOC_ALL = TOC_LEFT + TOC_RIGHT

FINAL_REMARKS_TITLE = "FINAL REMARKS: THE SOVEREIGN AWAKENING"
FINAL_REMARKS_PARAGRAPHS = [
    "     You have officially arrived at the edge of the architectural blueprints. By closing this text, you are not merely finishing a manuscript; you are actively deciding to stop consenting to a digital simulation designed to profit off your emotional fragmentation and financial anxiety. The universe has never been a static collection of cold matter. It is a living, responsive information matrix waiting for you to stabilize your consciousness, lock onto divine certainty, and rewrite the default social programming of your twenties and thirties.",
    "     True advancement is an entirely unified multi-dimensional sport. You cannot build a lasting financial empire while hosting a fragmented, scarcity-minded heart, nor can you experience a sacred, subatomic romance if your mind is constantly mourning the ghosts of parallel choices on a digital feed. When you confine your total faith to the divine, you drop the fragile checklist of modern dating apps and step into real cosmic resonance. You become quantumly entangled with the infinite, transforming money into a tool for sovereignty and your romantic partnership into a locked, illuminating orbit.",
    "     If you have read this far, a fundamental shift has already occurred within the deepest layers of your subconscious mind. Your perspective has been permanently altered, your old illusions have been shattered, and you are now fully equipped to begin the active process of reality engineering. If the journey itself has not already cracked open the matrix of how you see the world, look closer at the glitches in your daily routine. You are no longer the automated product of a generational script; you are the architect, standing at the command console, completely ready to reprogram your life from the absolute source code upward.",
    "     The Matrix wins the absolute second you choose to remain small, safe, and hyper-independent. Now is the exact moment to crash your local system. Grab the hand of the soul who refuses to run from your complexity, anchor your daily actions in absolute spiritual clarity, and walk boldly forward into your self-authored destiny. Wake up from the collective dream state, trust the underlying sacred geometry of your path, and go build an extraordinary reality that forces the entire simulation to glitch.",
]


def dotted_line(title: str, page: int, width: int = 68) -> str:
    base = f"{title} "
    dots = "." * max(4, width - len(base) - len(str(page)))
    return f"{base}{dots}{page}"


def build_markdown_toc() -> str:
    lines = [
        "## 1.4 Table of Contents",
        "",
        "| Canonical Index |",
        "| --- |",
    ]
    for title, page in TOC_ALL:
        if page is None:
            lines.append(f"| **{title}** |")
        else:
            lines.append(f"| {dotted_line(title, page, width=78)} |")
    return "\n".join(lines)


def update_master_markdown() -> None:
    text = SOURCE_MD.read_text(encoding="utf-8")
    if FINAL_REMARKS_TITLE in text:
        text = text[: text.index(f"# {FINAL_REMARKS_TITLE}")].rstrip() + "\n"

    toc_pattern = re.compile(
        r"## 1\.4 Table of Contents.*?(?=\n## 1\.5 About the Author)",
        flags=re.S,
    )
    text = toc_pattern.sub(build_markdown_toc() + "\n", text)

    text = text.rstrip() + "\n\n# " + FINAL_REMARKS_TITLE + "\n\n"
    text += "\n\n".join(FINAL_REMARKS_PARAGRAPHS) + "\n"
    MASTER_MD.write_text(text, encoding="utf-8", newline="\n")


def set_cell_borders_none(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "nil")


def set_paragraph_font(paragraph, size_pt: float = 11.5, bold: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = "Sabon"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Sabon")
        run.font.size = Pt(size_pt)
        run.bold = bold


def clear_cell(cell) -> None:
    for paragraph in list(cell.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)


def add_toc_line(cell, title: str, page: int | None) -> None:
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    if page is None:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(title)
        run.bold = True
        run.font.name = "Sabon"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Sabon")
        run.font.size = Pt(12)
        return

    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.05), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )
    run = paragraph.add_run(title)
    run.font.name = "Sabon"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Sabon")
    run.font.size = Pt(11.5)
    run = paragraph.add_run("\t" + str(page))
    run.font.name = "Sabon"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Sabon")
    run.font.size = Pt(11.5)


def update_toc_table(doc: Document) -> None:
    table = doc.tables[0]
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    merged = table.cell(0, 0).merge(table.cell(0, 1))
    clear_cell(merged)
    merged.width = Inches(6.5)
    set_cell_borders_none(merged)

    for title, page in TOC_ALL:
        add_toc_line(merged, title, page)


def add_final_remarks(doc: Document) -> None:
    existing = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    if FINAL_REMARKS_TITLE in existing:
        return

    doc.add_page_break()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(36)
    run = title.add_run(FINAL_REMARKS_TITLE)
    run.bold = True
    run.font.name = "Sabon"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Sabon")
    run.font.size = Pt(12)

    for text in FINAL_REMARKS_PARAGRAPHS:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(10)
        paragraph.paragraph_format.line_spacing = Pt(14)
        run = paragraph.add_run(text)
        run.font.name = "Sabon"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Sabon")
        run.font.size = Pt(11.5)


def main() -> None:
    if not BASE_DOCX.exists():
        raise FileNotFoundError(BASE_DOCX)
    if not AUTHOR_PHOTO.exists():
        raise FileNotFoundError(AUTHOR_PHOTO)

    update_master_markdown()
    shutil.copy2(BASE_DOCX, MASTER_DOCX)
    doc = Document(MASTER_DOCX)
    update_toc_table(doc)
    add_final_remarks(doc)
    doc.save(MASTER_DOCX)
    print("MASTER_DOCX", MASTER_DOCX)
    print("MASTER_MD", MASTER_MD)


if __name__ == "__main__":
    main()
